import os
import json
import re
import zipfile
from typing import List
from docx import Document
import gradio as gr
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from pydantic import BaseModel
from langchain_core.documents import Document as LCDocument

# ---------------- CONFIG ---------------- #
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise RuntimeError("GOOGLE_API_KEY not set in .env")

TOP_N_CHUNKS = 10  # configurable retrieval limit

# ---------------- MODELS ---------------- #
class Issue(BaseModel):
    document: str
    section: str
    issue: str
    severity: str
    suggestion: str

class AnalysisResult(BaseModel):
    issues_found: List[Issue]

# ---------------- AGENT ---------------- #
class CorporateAgent:
    def __init__(self, knowledge_base_path: str):
        print("Initializing Corporate Agent...")
        self.llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=GOOGLE_API_KEY)
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
        self.knowledge_base_path = knowledge_base_path
        print("✅ Corporate Agent is ready.")

    def _list_subfolders(self):
        subfolders = []
        for root, dirs, _ in os.walk(self.knowledge_base_path):
            for d in dirs:
                subfolders.append(os.path.join(root, d))
        return subfolders

    def _ask_ai_to_select_folders(self, doc_preview, folder_list):
        print("📁 All available subfolders:")
        for f in folder_list:
            print("   ", f)

        folder_names = "\n".join(folder_list)
        prompt = f"""
You are an ADGM legal assistant.
Preview of uploaded document:
---
{doc_preview}
---

Here are available reference folders:
{folder_names}

Select ONLY the most relevant folders for reviewing the document.
Be concise.
Return ONLY a JSON list of folder paths, e.g.: ["path1", "path2"]
"""
        print("🤖 Asking AI to select folders...")
        resp = self.llm.invoke(prompt)
        print("🤖 AI Folder Selection Response:")
        print(resp.content)
        resp_text = resp.content.replace("```json", "").replace("```", "").strip()
        try:
            selected = json.loads(resp_text)
            print("📂 AI selected folders (parsed):", selected)
            return [f for f in selected if os.path.exists(f)]
        except Exception as e:
            print("❌ Error parsing folder selection from AI:", e)
            print(f"   Raw AI response: {resp_text}")
            return []

    def _load_docs_from_folders(self, folder_paths):
        docs = []
        for folder in folder_paths:
            loaded_count = 0
            for root, _, files in os.walk(folder):
                for file in files:
                    path = os.path.join(root, file)
                    if file.lower().endswith(".docx"):
                        loader = Docx2txtLoader(path)
                    elif file.lower().endswith(".txt"):
                        loader = TextLoader(path, encoding="utf-8")
                    elif file.lower().endswith(".pdf"):
                        loader = PyMuPDFLoader(path)
                    else:
                        continue
                    try:
                        loaded = loader.load()
                        docs.extend(loaded)
                        loaded_count += len(loaded)
                    except Exception as e:
                        print(f"❌ Could not load {path}: {e}")
            print(f"   📄 Loaded {loaded_count} files from {folder}")
        return docs

    def _boost_relevance_by_filename(self, docs, uploaded_text):
        boosted_docs = []
        for d in docs:
            boost = 0
            fname = d.metadata.get("source", "").lower()
            if any(keyword.lower() in fname for keyword in uploaded_text.split()):
                boost = 0.1  # small similarity boost
            boosted_docs.append(LCDocument(page_content=d.page_content, metadata={**d.metadata, "boost": boost}))
        return boosted_docs

    def _retrieve_top_chunks(self, docs, uploaded_text):
        docs = self._boost_relevance_by_filename(docs, uploaded_text)
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        split_docs = splitter.split_documents(docs)
        vectorstore = FAISS.from_documents(split_docs, self.embeddings)
        retriever = vectorstore.as_retriever(search_kwargs={"k": TOP_N_CHUNKS})
        return retriever

    def _create_analysis_chain(self, retriever):
        system_prompt = """
You are an ADGM legal compliance expert.
Using ONLY the provided reference context, review the uploaded document for compliance issues.
Return JSON: {{ "issues_found": [ {{ "section": "..." }}, {{ "issue": "..." }}, {{ "severity": "..." }}, {{ "suggestion": "..." }} ] }}

<context>
{context}
</context>

Document Text:
{input}
"""
        prompt = ChatPromptTemplate.from_template(system_prompt)
        qa_chain = create_stuff_documents_chain(self.llm, prompt)
        return create_retrieval_chain(retriever, qa_chain)

    def _read_docx_text(self, file_path):
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    def _add_comments_to_docx(self, file_path, issues):
        doc = Document(file_path)
        for issue in issues:
            for para in doc.paragraphs:
                if issue["section"] in para.text:
                    comment_text = f"Issue: {issue['issue']}\nSeverity: {issue['severity']}\nSuggestion: {issue['suggestion']}"
                    para.text += f"  [COMMENT: {comment_text}]"
                    break
        return doc

    def safe_json_parse(self, ai_output):
        """Cleans AI output and parses it into JSON."""
        if hasattr(ai_output, "content"):
            raw_output = ai_output.content
        elif isinstance(ai_output, dict) and "answer" in ai_output:
            raw_output = ai_output["answer"]
        else:
            raw_output = str(ai_output)

        # Debug print
        print("\n🤖 Raw AI output before cleaning:\n", raw_output, "\n")

        # Remove markdown code fences and language tags
        clean_output = re.sub(r"```(json)?", "", raw_output).strip()
        clean_output = clean_output.replace("```", "").strip()

        # Extract first valid JSON object if extra text is present
        json_match = re.search(r"\{[\s\S]*\}", clean_output)
        if json_match:
            clean_output = json_match.group(0)

        # Debug print
        print("\n🧹 Cleaned AI output:\n", clean_output, "\n")

        try:
            return json.loads(clean_output)
        except json.JSONDecodeError as e:
            print("❌ JSON parsing failed:", e)
            return None


    def review_documents(self, files: list):
        if not files:
            return "Please upload at least one document.", {}, None

        all_issues = []
        reviewed_doc_paths = []

        for file_obj in files:
            doc_text = self._read_docx_text(file_obj.name)
            preview = "\n".join(doc_text.split("\n")[:10])

            subfolders = self._list_subfolders()
            selected_folders = self._ask_ai_to_select_folders(preview, subfolders)
            if not selected_folders:
                return "No relevant folders selected by AI.", {}, None

            reference_docs = self._load_docs_from_folders(selected_folders)
            if not reference_docs:
                return "No reference documents found in selected folders.", {}, None

            retriever = self._retrieve_top_chunks(reference_docs, doc_text)
            analysis_chain = self._create_analysis_chain(retriever)

            response = analysis_chain.invoke({"input": doc_text})
            print("🤖 AI Analysis Response (raw object):", response)

            # Safely parse AI response into JSON
            analysis_data = self.safe_json_parse(response)

            if not analysis_data:
                return "❌ Could not parse AI analysis into JSON.", {}, None

            issues = analysis_data.get("issues_found", [])
            for issue in issues:
                issue["document"] = os.path.basename(file_obj.name)
            all_issues.extend(issues)

            reviewed_doc = self._add_comments_to_docx(file_obj.name, issues)
            reviewed_path = f"reviewed_{os.path.basename(file_obj.name)}"
            reviewed_doc.save(reviewed_path)
            reviewed_doc_paths.append(reviewed_path)

        report = {
            "process": "Document Review",
            "documents_uploaded": len(files),
            "required_documents": None,
            "missing_document": "",
            "issues_found": all_issues
        }
        message = f"Review complete. Found {len(all_issues)} issues." if all_issues else "No issues found."

        zip_path = "reviewed_documents.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for file_path in reviewed_doc_paths:
                zip_file.write(file_path, os.path.basename(file_path))
        for file_path in reviewed_doc_paths:
            os.remove(file_path)

        return message, report, zip_path


# ---------------- MAIN APP ---------------- #
def main():
    knowledge_base_path = "ADGM_REFERENCES"
    os.makedirs(knowledge_base_path, exist_ok=True)
    agent = CorporateAgent(knowledge_base_path)

    def gradio_interface_fn(files):
        message, report_dict, zip_path = agent.review_documents(files)
        return message, report_dict, zip_path

    with gr.Blocks(title="ADGM Corporate Agent") as demo:
        with gr.Tab("Document Reviewer"):
            file_input = gr.File(label="Upload .docx Documents", file_count="multiple", file_types=[".docx"])
            submit_btn = gr.Button("Review Documents")
            output_message = gr.Textbox(label="Summary")
            output_json = gr.JSON(label="Analysis Report")
            output_download = gr.File(label="Download Reviewed Documents")
            submit_btn.click(
                fn=gradio_interface_fn,
                inputs=[file_input],
                outputs=[output_message, output_json, output_download]
            )

    demo.launch()

if __name__ == "__main__":
    main()
